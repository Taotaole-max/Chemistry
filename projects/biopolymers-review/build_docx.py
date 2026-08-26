"""Build the CM3254 biopolymer review report.md into a Word document.

Purpose-built for this report's specific markdown structure (headings, one
title, one author line, one blockquote, markdown tables, figure placeholders,
bullet lists, bold/italic/code inline spans, HTML comments). Not a general
markdown-to-docx converter.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

BASE = Path(__file__).parent
SRC = Path(sys.argv[1]) if len(sys.argv) > 1 else BASE / "report.md"
FIG_DIR = BASE / "figures" / "output"
OUT = Path(sys.argv[2]) if len(sys.argv) > 2 else BASE / "Biopolymers-review-DRAFT.docx"
EAST_ASIA_FONT = sys.argv[3] if len(sys.argv) > 3 else "Times New Roman"

FIGURE_FILES = {
    1: "fig1_classification.png",
    2: "fig7_causality_chain.png",
    3: "fig_polysaccharide_monomers.png",
    4: "fig3_cellulose_hierarchy.png",
    5: "fig_polysaccharide_chains.png",
    6: "fig_polyester_structures.png",
    7: "fig_protein_structures.png",
    8: "fig_lignin_monomers.png",
    9: "fig_rubber_monomer.png",
    10: "fig9_dispersity.png",
    11: "fig4_thermal_windows.png",
    12: "fig5_property_map.png",
    13: "fig6a_polyester_hydrolysis.png",
    14: "fig6b_enzymatic_degradation.png",
    15: "fig6c_hydrolysis_barriers.png",
    16: "fig10_processing_routes.png",
    17: "fig8_nr_sbr_comparison.png",
    18: "fig12_appendix_monomers.png",
    19: "fig13_appendix_chain_structures.png",
}

# Page-budget control: never insert a figure at native/full text width — cap its
# *height* instead, and derive width from the image's own aspect ratio.
#
# Unified 3-band system (2026-08-26 pass, replacing the earlier ad hoc per-figure
# values that ranged 1.55-4.2in with no consistent logic and looked inconsistent in
# print): S = simple/small single-panel figure, M = standard content figure,
# L = dense multi-panel or multi-row figure. Every BODY figure (1-17) takes one of
# exactly these three values. Appendix figures (18-19) are exempt from the body's
# 10-page budget (OUTLINE.md), so they may exceed L where legibility calls for it —
# F19 is the one documented exception, a genuinely dense 6-panel schematic.
FIGURE_HEIGHT_S = 1.3
FIGURE_HEIGHT_M = 1.9
FIGURE_HEIGHT_L = 2.3
FIGURE_MAX_HEIGHT_IN_DEFAULT = FIGURE_HEIGHT_M
FIGURE_MAX_HEIGHT_IN = {
    1: FIGURE_HEIGHT_L,   # classification tree + origin-tag legend, multi-row
    2: FIGURE_HEIGHT_L,   # 7-row causality-chain table
    3: FIGURE_HEIGHT_S,   # polysaccharide monomer strip, wide/short
    4: FIGURE_HEIGHT_M,   # cellulose chain->sheet->crystal, 3 panels
    5: FIGURE_HEIGHT_M,   # amylose/alginate/chitosan chain structures, 3 panels
    6: FIGURE_HEIGHT_S,   # PHB/PLLA monomers + helix packing, 3 panels
    7: FIGURE_HEIGHT_M,   # silk/collagen monomers + 2° structure, 2x2
    8: FIGURE_HEIGHT_S,   # 3 monolignols
    9: FIGURE_HEIGHT_S,   # single small monomer
    10: FIGURE_HEIGHT_M,  # 6-row dispersity comparison with annotations
    11: FIGURE_HEIGHT_M,  # thermal windows, one row per material
    12: FIGURE_HEIGHT_M,  # Ashby-style mechanical map
    13: FIGURE_HEIGHT_S,  # single hydrolysis reaction scheme
    14: FIGURE_HEIGHT_S,  # crystalline/amorphous accessibility bar
    15: FIGURE_HEIGHT_M,  # energy-barrier curves + axes
    16: FIGURE_HEIGHT_M,  # processing decision flowchart
    17: FIGURE_HEIGHT_S,  # NR/SBR chain-scale structure, 4 small panels
    18: FIGURE_HEIGHT_L,  # appendix: 10-panel monomer gallery, wide and short
    19: 4.2,               # appendix: 6-panel chain-structure schematic, tall —
                            # documented exception, see comment above
}
FIGURE_MAX_WIDTH_IN = 6.2

# Table-as-image placeholder: "*[TableImage — <filename>]*" placed right after a
# "**Table N.**"/"**表 N.**" caption, in place of the markdown pipe table. The filename
# is literal (not looked up by number) so report.md and report_zh.md can each point at
# their own language-specific render of the same data.
TABLE_IMAGE_MAX_HEIGHT_IN = 3.4


def strip_html_comments(text: str) -> str:
    return re.sub(r"<!--.*?-->", "", text, flags=re.S)


def set_cell_borders(cell, top=False, bottom=False, size=8):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for side, on in (("top", top), ("bottom", bottom)):
        el = borders.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            borders.append(el)
        if on:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(size))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "nil")


def add_inline_runs(paragraph, text, base_italic=False, base_bold=False):
    """Parse **bold**, *italic*, `code` spans (non-nested) into runs."""
    pattern = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.italic = base_italic
        elif part.startswith("*") and part.endswith("*") and len(part) > 1:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.italic = True
        else:
            run = paragraph.add_run(part)
            run.italic = base_italic
            run.bold = base_bold


def is_special_line_start(s):
    """True if a stripped line begins a new block element (not a wrapped continuation
    of the previous paragraph/list-item/reference-entry)."""
    if s == "" or s == "---":
        return True
    if s.startswith("# "):
        return True
    if s.startswith("*") and s.endswith("*") and not s.startswith("**") and not s.startswith("*["):
        return True
    if s.startswith(">"):
        return True
    if re.match(r"^\*\[(Figure|图) \d+ — ", s):
        return True
    if re.match(r"^\*\[TableImage — ", s):
        return True
    if re.match(r"^#{2,4}\s+", s):
        return True
    if s.startswith("|"):
        return True
    if s.startswith("- "):
        return True
    if re.match(r"^\*\*(Table|表) \d+\.\*\*", s):
        return True
    if re.match(r"^\d{1,3}\.\s", s):  # numbered reference-list entry
        return True
    return False


def gather_paragraph(lines, i, n):
    """Join the current line and any following wrapped-continuation lines (source
    hard-wraps a single logical paragraph/list-item/reference across several lines;
    only a blank line or another block-start line ends it) into one string."""
    buf = [lines[i].strip()]
    i += 1
    while i < n:
        nxt = lines[i].strip()
        if is_special_line_start(nxt):
            break
        buf.append(nxt)
        i += 1
    return " ".join(buf), i


def parse_md_table(lines):
    """lines: list of '|...|' rows including the header-separator row. Returns list of row-cell-lists."""
    rows = []
    for i, ln in enumerate(lines):
        if i == 1:
            continue  # the |---|---| separator row
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def build():
    text = strip_html_comments(SRC.read_text(encoding="utf-8"))
    raw_lines = text.split("\n")

    # Merge wrapped "*[Figure N — ...]*" captions that span multiple source lines
    # back into a single logical line before the main parse loop.
    lines = []
    buf = None
    for ln in raw_lines:
        if buf is not None:
            buf += " " + ln.strip()
            if buf.rstrip().endswith("]*"):
                lines.append(buf)
                buf = None
            continue
        if re.match(r"^\s*\*\[(?:Figure|图) \d+ — ", ln) and not ln.strip().endswith("]*"):
            buf = ln.strip()
            continue
        lines.append(ln)
    if buf is not None:
        lines.append(buf)

    doc = Document()

    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(2)

    try:
        list_bullet = doc.styles["List Bullet"]
        list_bullet.font.name = "Times New Roman"
        list_bullet.font.size = Pt(12)
        list_bullet.paragraph_format.line_spacing = 1.5
        list_bullet.paragraph_format.space_after = Pt(3)
    except KeyError:
        pass
    rpr = normal.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)

    for i in range(1, 5):
        try:
            h = doc.styles[f"Heading {i}"]
            h.font.name = "Times New Roman"
            h.font.color.rgb = None
            h.font.size = Pt({1: 14, 2: 13, 3: 12, 4: 12}[i])
            h.font.bold = True
            h.paragraph_format.space_before = Pt(6)
            h.paragraph_format.space_after = Pt(2)
            h.paragraph_format.line_spacing = 1.0
            h.paragraph_format.keep_with_next = False
            h.paragraph_format.keep_together = False
            h.paragraph_format.widow_control = True
            h_rpr = h.element.get_or_add_rPr()
            h_rFonts = h_rpr.find(qn("w:rFonts"))
            if h_rFonts is None:
                h_rFonts = OxmlElement("w:rFonts")
                h_rpr.append(h_rFonts)
            h_rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        except KeyError:
            pass

    n = len(lines)
    i = 0
    first_para_done = False
    in_references = False
    fig_used = set()

    while i < n:
        raw = lines[i]
        s = raw.strip()

        if s == "":
            i += 1
            continue

        if s == "---":
            i += 1
            continue

        # Title (first # line)
        if s.startswith("# ") and not first_para_done:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(s[2:].strip())
            run.bold = True
            run.font.size = Pt(16)
            first_para_done = True
            i += 1
            continue

        # Author / italic byline line: "*...*"
        if s.startswith("*") and s.endswith("*") and not s.startswith("**") and not s.startswith("*["):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(p, s)
            i += 1
            continue

        # Blockquote (AI declaration / translator's note) — join hard-wrapped ">" lines
        # back into one logical paragraph, same reasoning as gather_paragraph() below.
        if s.startswith(">"):
            parts = [s.lstrip(">").strip()]
            i += 1
            while i < n and lines[i].strip().startswith(">") and lines[i].strip().lstrip(">").strip():
                parts.append(lines[i].strip().lstrip(">").strip())
                i += 1
            p = doc.add_paragraph()
            add_inline_runs(p, " ".join(parts), base_italic=True)
            continue

        # Figure placeholder line: *[Figure N — ...]* or *[图 N — ...]*
        m = re.match(r"^\*\[(Figure|图) (\d+) — (.+)\]\*$", s)
        if m:
            label = m.group(1)
            fig_num = int(m.group(2))
            caption_text = m.group(3)
            fname = FIGURE_FILES.get(fig_num)
            fpath = FIG_DIR / fname if fname else None
            if fpath and fpath.exists():
                native_w, native_h = Image.open(fpath).size
                aspect_h_over_w = native_h / native_w
                max_h = FIGURE_MAX_HEIGHT_IN.get(fig_num, FIGURE_MAX_HEIGHT_IN_DEFAULT)
                width_in = min(FIGURE_MAX_WIDTH_IN, max_h / aspect_h_over_w)
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_after = Pt(2)
                # Keep the image glued to its caption paragraph — without this, Word can
                # push just the caption to the next page, leaving an orphaned caption line
                # on an otherwise-blank page (found via PDF visual QA on appendix Fig 19).
                p_img.paragraph_format.keep_with_next = True
                run = p_img.add_run()
                run.add_picture(str(fpath), width=Inches(width_in))
                fig_used.add(fig_num)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(4)
            sep = ". " if label == "Figure" else "．"
            run = p.add_run(f"{label} {fig_num}{sep}{caption_text}")
            run.italic = True
            run.font.size = Pt(10)
            i += 1
            continue

        # Table-as-image placeholder: *[TableImage — <filename>]*, sits right after a
        # "**Table N.**"/"**表 N.**" caption in place of the markdown pipe table.
        m = re.match(r"^\*\[TableImage — (.+)\]\*$", s)
        if m:
            fpath = FIG_DIR / m.group(1).strip()
            if fpath.exists():
                native_w, native_h = Image.open(fpath).size
                aspect_h_over_w = native_h / native_w
                width_in = min(FIGURE_MAX_WIDTH_IN, TABLE_IMAGE_MAX_HEIGHT_IN / aspect_h_over_w)
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_after = Pt(4)
                run = p_img.add_run()
                run.add_picture(str(fpath), width=Inches(width_in))
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{2,4})\s+(.*)$", s)
        if m:
            level = len(m.group(1)) - 1
            heading_text = m.group(2)
            doc.add_heading(heading_text, level=level)
            if heading_text.strip().startswith("References"):
                in_references = True
            i += 1
            continue

        # Markdown table (possibly preceded by a bold "**Table N.** caption" paragraph already handled as normal para)
        if s.startswith("|"):
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = parse_md_table(table_lines)
            if rows:
                ncols = len(rows[0])
                table = doc.add_table(rows=0, cols=ncols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False
                table.allow_autofit = False

                # Proportional column widths (by max content length per column) instead of
                # Word's default autofit, which tends to under-pack long text columns and
                # bloat row height with unnecessary wraps. Raw character-count proportions
                # starve short-word columns (e.g. a "Polysaccharide" header column next to a
                # long free-text column) into wrapping a single word — use sqrt-weighting to
                # compress that spread, then floor each column at its own longest word so no
                # single unsplittable word is ever narrower than its column.
                col_max_len = [
                    max(len(r[c]) if c < len(r) else 0 for r in rows) for c in range(ncols)
                ]
                col_max_len = [max(v, 4) for v in col_max_len]
                col_weight = [v ** 0.5 for v in col_max_len]
                col_word_len = [
                    max((len(w) for r in rows if c < len(r) for w in r[c].split()), default=4)
                    for c in range(ncols)
                ]
                table_width_in = 6.2
                min_col_in = 0.45
                per_char_in = 0.078  # ~ Arial 8pt average glyph width
                col_floor_in = [max(min_col_in, wl * per_char_in) for wl in col_word_len]
                extra_in = max(0.0, table_width_in - sum(col_floor_in))
                total_weight = sum(col_weight)
                col_widths_in = [f + extra_in * w / total_weight
                                 for f, w in zip(col_floor_in, col_weight)]
                if sum(col_widths_in) > table_width_in:  # floors alone exceeded budget
                    scale = table_width_in / sum(col_widths_in)
                    col_widths_in = [w * scale for w in col_widths_in]

                tblPr = table._tbl.tblPr
                tblLayout = OxmlElement("w:tblLayout")
                tblLayout.set(qn("w:type"), "fixed")
                tblPr.append(tblLayout)

                grid = table._tbl.tblGrid
                for gridCol, w_in in zip(grid.findall(qn("w:gridCol")), col_widths_in):
                    gridCol.set(qn("w:w"), str(int(w_in * 1440)))

                for r_idx, row_cells in enumerate(rows):
                    row = table.add_row()
                    for c_idx, cell_text in enumerate(row_cells):
                        if c_idx >= ncols:
                            continue
                        cell = row.cells[c_idx]
                        cell.width = Inches(col_widths_in[c_idx])
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_after = Pt(0.5)
                        p.paragraph_format.space_before = Pt(0.5)
                        p.paragraph_format.line_spacing = 1.0
                        add_inline_runs(p, cell_text)
                        for run in p.runs:
                            run.font.size = Pt(8)
                            run.font.name = "Arial"
                            if r_idx == 0:
                                run.bold = True
                        set_cell_borders(
                            cell,
                            top=(r_idx == 0),
                            bottom=(r_idx == 0 or r_idx == len(rows) - 1),
                        )
                sp = doc.add_paragraph()
                sp.paragraph_format.space_after = Pt(0)
                sp.paragraph_format.line_spacing = 1.0
            continue

        # Bullet list item (source hard-wraps long items across several lines)
        if s.startswith("- "):
            text, i = gather_paragraph(lines, i, n)
            text = text[2:].strip() if text.startswith("- ") else text
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, text)
            continue

        # Table caption line, e.g. "**Table 2.** Provisional summary ..." — single-spaced,
        # smaller, tighter, like a figure caption. Also hard-wrapped across lines in source.
        if re.match(r"^\*\*(Table|表) \d+\.\*\*", s):
            text, i = gather_paragraph(lines, i, n)
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(2)
            add_inline_runs(p, text)
            for run in p.runs:
                run.font.size = Pt(10)
            continue

        # Regular paragraph (justified body text) — join hard-wrapped source lines first.
        text, i = gather_paragraph(lines, i, n)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if in_references:
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(4)
        add_inline_runs(p, text)

    doc.save(str(OUT))
    print("Saved:", OUT)
    print("Figures embedded:", sorted(fig_used), "/ expected", sorted(FIGURE_FILES))


if __name__ == "__main__":
    build()
